"""Convert report.md to a properly formatted Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD = "report.md"
OUT = "report.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Styles ────────────────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level, size, bold=True, center=False, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, bold=bold, size=size)
    return p

def body(text, italic=False, space_before=0, space_after=6, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = alignment
    # Handle inline bold/italic
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True, size=11)
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            set_font(run, italic=True, size=11)
        else:
            run = p.add_run(part)
            set_font(run, italic=italic, size=11)
    return p

def add_table(rows_data, header=True):
    nrows = len(rows_data)
    ncols = len(rows_data[0])
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = cell_text.strip()
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    if i == 0 and header:
                        run.bold = True
    # shade header row
    if header:
        for cell in tbl.rows[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D9D9D9')
            tcPr.append(shd)
    doc.add_paragraph()  # spacing after table

# ─────────────────────────────────────────────────────────────────────────────
# Parse and render the markdown
# ─────────────────────────────────────────────────────────────────────────────
with open(MD, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_table = False
table_rows = []

def flush_table():
    global table_rows, in_table
    if table_rows:
        # Remove separator rows (---|---|---)
        clean = [r for r in table_rows if not all(c.strip().replace('-','').replace(':','') == '' for c in r)]
        if clean:
            add_table(clean)
    table_rows = []
    in_table = False

while i < len(lines):
    line = lines[i].rstrip('\n')

    # Table row
    if line.startswith('|'):
        in_table = True
        cells = [c for c in line.split('|')]
        # Remove leading/trailing empty cells from split
        cells = cells[1:-1] if cells[0] == '' else cells
        table_rows.append(cells)
        i += 1
        continue

    # End of table
    if in_table and not line.startswith('|'):
        flush_table()

    # Skip blank lines (but add spacing)
    if line.strip() == '':
        i += 1
        continue

    # Title (# )
    if re.match(r'^# [^#]', line):
        text = line[2:].strip()
        heading(text, 1, size=16, center=True, space_before=0, space_after=4)

    # Author / metadata lines under title
    elif line.startswith('**Author:**') or line.startswith('**Course:**') or line.startswith('**Date:**') or line.startswith('**Affiliation:**'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(re.sub(r'\*\*([^*]+)\*\*', r'\1', line))
        set_font(run, size=11)

    # Horizontal rule
    elif line.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        p.add_run('─' * 80)

    # ## heading
    elif re.match(r'^## [^#]', line):
        text = line[3:].strip()
        heading(text, 2, size=13, space_before=14, space_after=4)

    # ### heading
    elif re.match(r'^### [^#]', line):
        text = line[4:].strip()
        heading(text, 3, size=12, space_before=10, space_after=3)

    # Image  ![alt](path)
    elif line.startswith('!['):
        m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
        if m:
            img_path = m.group(2)
            try:
                doc.add_picture(img_path, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                body(f'[Figure: {img_path}]', italic=True)

    # Bullet list (- item)
    elif re.match(r'^- ', line):
        text = line[2:].strip()
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_font(run, bold=True, size=11)
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                set_font(run, italic=True, size=11)
            else:
                run = p.add_run(part)
                set_font(run, size=11)

    # Block quote / note (starts with >)
    elif line.startswith('>'):
        text = line[1:].strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        set_font(run, italic=True, size=11)

    # Italic note line (starts with *)
    elif line.startswith('*Note:') or (line.startswith('*') and line.endswith('*')):
        text = line.strip('*').strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        set_font(run, italic=True, size=10)

    # Math / formula line
    elif line.strip().startswith('$$') or 'mathcal' in line:
        raw = line.strip().strip('$').strip()
        # Convert common LaTeX to readable Unicode plaintext
        for latex, plain in [
            (r'\mathcal{L}_{\text{src}}', 'L_src'),
            (r'\mathcal{L}_{\text{tgt}}', 'L_tgt'),
            (r'\mathcal{L}', 'L'),
            (r'\lambda_u', 'λ_u'),
            (r'\cdot', '·'),
            (r'\times', '×'),
            (r'\leq', '≤'),
            (r'\geq', '≥'),
            (r'\approx', '≈'),
        ]:
            raw = raw.replace(latex, plain)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(raw)
        set_font(run, italic=True, size=11)

    # Regular paragraph
    else:
        body(line)

    i += 1

# Flush any remaining table
if in_table:
    flush_table()

doc.save(OUT)
print(f"Saved: {OUT}")
